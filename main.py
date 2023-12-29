from docx import Document
from docx.shared import Pt
class main():
    def __init__(self ,headSize = 29 , headFont = "monospace" , normalSize = 15 , normalFont = "sans serif" , inFile = "in.txt" , outFile = "out.docx"):
        self.headSize = headSize
        self.headFont = headFont
        self.normalSize = normalSize
        self.normalFont = normalFont
        self.document =Document()

        self.inFile = inFile
        self.outfile = outFile
        self.points = []
        self.headers = []

        self.headStyle = self.document.styles["Title"]
        self.headStyle.font.name = self.headFont
        self.headStyle.font.size = Pt(self.headSize)

        self.normalStyle = self.document.styles["Normal"]
        self.normalStyle.font.name = self.normalFont
        self.normalStyle.font.size = Pt(self.normalSize)

    def readfile(self):
        file = open(self.inFile , "rt")
        self.fileContents = file.read()
        file.close()
    def get_info(self):
        for line in self.fileContents.splitlines():
            if line.startswith("!!HEAD"):
                self.headers.append(line)
            else:
                self.points.append(line)
    def Write(self):
        for line in self.fileContents.splitlines():
            print(line)
            if line.startswith("!!HEAD"):
                head = self.document.add_heading(text = line.replace("!!HEAD" , ""))
                head.style = self.document.styles["Title"]
            elif line.startswith("!!IMAGE"):
                self.document.add_picture(line.replace("!!IMAGE" , ""))
            else:
                paragraph = self.document.add_paragraph(text = line)
                paragraph.style = self.document.styles["Normal"]
        

    def save(self):
        self.document.save(self.outfile)

instance = main()
instance.readfile()
instance.Write()
instance.save()
